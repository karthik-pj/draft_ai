# app.py
import streamlit as st
import sqlite3
import hashlib
import os
from datetime import datetime
import json
import PyPDF2
import docx
import fitz  # PyMuPDF
import io
import re
import pandas as pd
import base64
import groq
import pinecone
from sentence_transformers import SentenceTransformer
import uuid

# =============================================================================
# CONFIGURATION AND SETUP
# =============================================================================

# Page configuration
st.set_page_config(
    page_title="AI Legal Draft Assistant",
    page_icon="⚖️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# =============================================================================
# AUTHENTICATION MODULE
# =============================================================================

DB_NAME = "legal_assistant.db"

def get_db_connection():
    """Get database connection with proper error handling"""
    try:
        conn = sqlite3.connect(DB_NAME)
        conn.row_factory = sqlite3.Row
        return conn
    except sqlite3.Error as e:
        st.error(f"Database connection error: {e}")
        return None

def init_db():
    """Initialize SQLite database with proper error handling"""
    conn = get_db_connection()
    if conn is None:
        return False
    
    cursor = conn.cursor()
    
    try:
        # Users table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                name TEXT NOT NULL,
                email TEXT NOT NULL,
                role TEXT NOT NULL,
                active BOOLEAN DEFAULT 1,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                created_by TEXT,
                last_login TIMESTAMP
            )
        ''')
        
        # Audit log table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS audit_log (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                username TEXT,
                action TEXT,
                details TEXT,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')
        
        # Drafts table to store all drafts in database
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS drafts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                username TEXT,
                case_type TEXT,
                draft_content TEXT,
                medical_text TEXT,
                case_details TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        ''')
        
        # Training documents table
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS training_docs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                document_name TEXT,
                document_type TEXT,
                content TEXT,
                uploaded_by TEXT,
                uploaded_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        # Insert default admin user if not exists
        cursor.execute("SELECT COUNT(*) FROM users WHERE username = 'admin'")
        if cursor.fetchone()[0] == 0:
            admin_password_hash = hash_password("admin123")
            cursor.execute('''
                INSERT INTO users (username, password_hash, name, email, role, created_by)
                VALUES (?, ?, ?, ?, ?, ?)
            ''', ('admin', admin_password_hash, 'System Administrator', 'admin@legalfirm.com', 'admin', 'system'))
        
        conn.commit()
        return True
        
    except sqlite3.Error as e:
        st.error(f"Database initialization error: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def hash_password(password):
    """Hash password using SHA-256"""
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password, password_hash):
    """Verify password against hash"""
    return hash_password(password) == password_hash

def initialize_authentication():
    """Initialize session state for authentication"""
    if 'authenticated' not in st.session_state:
        st.session_state.authenticated = False
    if 'username' not in st.session_state:
        st.session_state.username = None
    if 'user_role' not in st.session_state:
        st.session_state.user_role = None
    if 'name' not in st.session_state:
        st.session_state.name = None
    if 'user_id' not in st.session_state:
        st.session_state.user_id = None
    
    # Initialize database
    if not init_db():
        st.error("Failed to initialize database. Please check the application logs.")

def login():
    """Login form with SQLite authentication"""
    with st.form("login_form"):
        st.subheader("Login to Legal Draft Assistant")
        
        username = st.text_input("Username", placeholder="Enter your username")
        password = st.text_input("Password", type="password", placeholder="Enter your password")
        
        login_button = st.form_submit_button("Login")
        
        if login_button:
            if not username or not password:
                st.error("Please enter both username and password")
            else:
                user = authenticate_user(username, password)
                if user:
                    st.session_state.authenticated = True
                    st.session_state.username = user['username']
                    st.session_state.user_role = user['role']
                    st.session_state.name = user['name']
                    st.session_state.user_id = user['id']
                    
                    # Update last login
                    update_last_login(user['id'])
                    
                    # Log login action
                    log_audit_action(user['id'], user['username'], "LOGIN", "User logged in successfully")
                    
                    # Load user's draft history from database
                    load_user_drafts(user['id'])
                    
                    st.success(f"Welcome {user['name']}!")
                    st.rerun()
                else:
                    st.error("Invalid username or password")
    
    return st.session_state.authenticated

def authenticate_user(username, password):
    """Authenticate user against database"""
    conn = get_db_connection()
    if conn is None:
        return None
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT id, username, password_hash, name, role 
            FROM users 
            WHERE username = ? AND active = 1
        ''', (username,))
        
        user = cursor.fetchone()
        if user and verify_password(password, user['password_hash']):
            return {
                'id': user['id'],
                'username': user['username'],
                'name': user['name'],
                'role': user['role']
            }
        return None
    except sqlite3.Error as e:
        st.error(f"Authentication error: {e}")
        return None
    finally:
        conn.close()

def update_last_login(user_id):
    """Update user's last login timestamp"""
    conn = get_db_connection()
    if conn is None:
        return
    
    try:
        cursor = conn.cursor()
        cursor.execute('UPDATE users SET last_login = CURRENT_TIMESTAMP WHERE id = ?', (user_id,))
        conn.commit()
    except sqlite3.Error as e:
        st.error(f"Error updating last login: {e}")
        conn.rollback()
    finally:
        conn.close()

def create_user(username, password, name, email, role, created_by):
    """Create a new user in database"""
    conn = get_db_connection()
    if conn is None:
        return False, "Database connection failed"
    
    try:
        cursor = conn.cursor()
        password_hash = hash_password(password)
        
        cursor.execute('''
            INSERT INTO users (username, password_hash, name, email, role, created_by)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (username, password_hash, name, email, role, created_by))
        
        user_id = cursor.lastrowid
        conn.commit()
        
        # Log the action
        creator_username = get_username_by_id(created_by)
        log_audit_action(created_by, creator_username, "CREATE_USER", f"Created user: {username}")
        
        return True, "User created successfully"
    except sqlite3.IntegrityError:
        return False, "Username already exists"
    except sqlite3.Error as e:
        return False, f"Database error: {str(e)}"
    except Exception as e:
        return False, f"Error creating user: {str(e)}"
    finally:
        conn.close()

def update_user(user_id, **kwargs):
    """Update user information with proper error handling"""
    if not user_id:
        return False, "Invalid user ID"
    
    conn = get_db_connection()
    if conn is None:
        return False, "Database connection failed"
    
    try:
        cursor = conn.cursor()
        update_fields = []
        update_values = []
        
        for key, value in kwargs.items():
            if key == 'password' and value:
                update_fields.append("password_hash = ?")
                update_values.append(hash_password(value))
            elif key in ['name', 'email', 'role'] and value is not None:
                update_fields.append(f"{key} = ?")
                update_values.append(value)
        
        if not update_fields:
            return False, "No valid fields to update"
        
        update_values.append(user_id)
        query = f"UPDATE users SET {', '.join(update_fields)} WHERE id = ?"
        cursor.execute(query, update_values)
        
        affected_rows = cursor.rowcount
        conn.commit()
        
        if affected_rows == 0:
            return False, "User not found or no changes made"
        
        # Log the action
        log_audit_action(st.session_state.user_id, st.session_state.username, "UPDATE_USER", f"Updated user ID: {user_id}")
        
        return True, "User updated successfully"
    except sqlite3.Error as e:
        conn.rollback()
        return False, f"Database error: {str(e)}"
    except Exception as e:
        conn.rollback()
        return False, f"Error updating user: {str(e)}"
    finally:
        conn.close()

def delete_user(user_id, username):
    """Soft delete a user with proper error handling"""
    if not user_id:
        return False, "Invalid user ID"
    
    if username == "admin":
        return False, "Cannot delete admin user"
    
    conn = get_db_connection()
    if conn is None:
        return False, "Database connection failed"
    
    try:
        cursor = conn.cursor()
        
        # First, verify the user exists and is active
        cursor.execute('SELECT username FROM users WHERE id = ? AND active = 1', (user_id,))
        user_exists = cursor.fetchone()
        
        if not user_exists:
            return False, "User not found or already deleted"
        
        # Perform the soft delete
        cursor.execute('UPDATE users SET active = 0 WHERE id = ?', (user_id,))
        affected_rows = cursor.rowcount
        conn.commit()
        
        if affected_rows == 0:
            return False, "No user was deleted"
        
        # Log the action
        log_audit_action(st.session_state.user_id, st.session_state.username, "DELETE_USER", f"Deleted user: {username} (ID: {user_id})")
        
        return True, f"User '{username}' deleted successfully"
    except sqlite3.Error as e:
        conn.rollback()
        return False, f"Database error: {str(e)}"
    except Exception as e:
        conn.rollback()
        return False, f"Error deleting user: {str(e)}"
    finally:
        conn.close()

def get_all_users():
    """Get all active users with proper error handling"""
    conn = get_db_connection()
    if conn is None:
        return {}
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT id, username, name, email, role, created_at, created_by, last_login
            FROM users 
            WHERE active = 1
            ORDER BY username
        ''')
        
        users = cursor.fetchall()
        user_dict = {}
        
        for user in users:
            user_dict[user['username']] = {
                'id': user['id'],
                'name': user['name'],
                'email': user['email'],
                'role': user['role'],
                'created_at': user['created_at'],
                'created_by': user['created_by'],
                'last_login': user['last_login']
            }
        
        return user_dict
    except sqlite3.Error as e:
        st.error(f"Error fetching users: {e}")
        return {}
    finally:
        conn.close()

def get_user_by_username(username):
    """Get user by username"""
    conn = get_db_connection()
    if conn is None:
        return None
    
    try:
        cursor = conn.cursor()
        cursor.execute('SELECT id, username, name, role FROM users WHERE username = ? AND active = 1', (username,))
        user = cursor.fetchone()
        
        if user:
            return {
                'id': user['id'],
                'username': user['username'],
                'name': user['name'],
                'role': user['role']
            }
        return None
    except sqlite3.Error as e:
        st.error(f"Error fetching user: {e}")
        return None
    finally:
        conn.close()

def get_username_by_id(user_id):
    """Get username by user ID"""
    conn = get_db_connection()
    if conn is None:
        return "Unknown"
    
    try:
        cursor = conn.cursor()
        cursor.execute('SELECT username FROM users WHERE id = ?', (user_id,))
        result = cursor.fetchone()
        return result['username'] if result else "Unknown"
    except sqlite3.Error:
        return "Unknown"
    finally:
        conn.close()

def log_audit_action(user_id, username, action, details):
    """Log user actions to audit trail"""
    conn = get_db_connection()
    if conn is None:
        return
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO audit_log (user_id, username, action, details)
            VALUES (?, ?, ?, ?)
        ''', (user_id, username, action, details))
        conn.commit()
    except sqlite3.Error as e:
        st.error(f"Error logging audit action: {e}")
        conn.rollback()
    finally:
        conn.close()

def get_audit_logs(limit=100):
    """Get audit logs"""
    conn = get_db_connection()
    if conn is None:
        return []
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT al.timestamp, al.username, al.action, al.details, u.name
            FROM audit_log al
            LEFT JOIN users u ON al.user_id = u.id
            ORDER BY al.timestamp DESC
            LIMIT ?
        ''', (limit,))
        
        return cursor.fetchall()
    except sqlite3.Error as e:
        st.error(f"Error fetching audit logs: {e}")
        return []
    finally:
        conn.close()

def save_draft_to_db(user_id, username, case_type, draft_content, medical_text, case_details):
    """Save draft to database"""
    conn = get_db_connection()
    if conn is None:
        return None
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO drafts (user_id, username, case_type, draft_content, medical_text, case_details)
            VALUES (?, ?, ?, ?, ?, ?)
        ''', (user_id, username, case_type, draft_content, medical_text, json.dumps(case_details)))
        
        draft_id = cursor.lastrowid
        conn.commit()
        
        # Log the action
        log_audit_action(user_id, username, "SAVE_DRAFT", f"Saved draft for case: {case_type}")
        
        return draft_id
    except Exception as e:
        st.error(f"Error saving draft: {e}")
        conn.rollback()
        return None
    finally:
        conn.close()

def load_user_drafts(user_id):
    """Load user's drafts from database into session state"""
    conn = get_db_connection()
    if conn is None:
        st.session_state.draft_history = []
        return
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT case_type, draft_content, created_at 
            FROM drafts 
            WHERE user_id = ? 
            ORDER BY created_at DESC
            LIMIT 50
        ''', (user_id,))
        
        drafts = cursor.fetchall()
        
        # Convert to session state format
        st.session_state.draft_history = []
        for draft in drafts:
            st.session_state.draft_history.append({
                'timestamp': draft['created_at'],
                'case_type': draft['case_type'],
                'user': st.session_state.username,
                'content': draft['draft_content']
            })
    except sqlite3.Error as e:
        st.error(f"Error loading drafts: {e}")
        st.session_state.draft_history = []
    finally:
        conn.close()

def get_user_drafts_count(user_id):
    """Get count of drafts for a user"""
    conn = get_db_connection()
    if conn is None:
        return 0
    
    try:
        cursor = conn.cursor()
        cursor.execute('SELECT COUNT(*) FROM drafts WHERE user_id = ?', (user_id,))
        count = cursor.fetchone()[0]
        return count
    except sqlite3.Error:
        return 0
    finally:
        conn.close()

def get_all_drafts_stats():
    """Get statistics about all drafts"""
    conn = get_db_connection()
    if conn is None:
        return {'total_drafts': 0, 'drafts_by_type': {}, 'recent_drafts': []}
    
    try:
        cursor = conn.cursor()
        
        # Total drafts count
        cursor.execute('SELECT COUNT(*) FROM drafts')
        total_drafts = cursor.fetchone()[0]
        
        # Drafts by case type
        cursor.execute('''
            SELECT case_type, COUNT(*) 
            FROM drafts 
            GROUP BY case_type 
            ORDER BY COUNT(*) DESC
        ''')
        drafts_by_type = dict(cursor.fetchall())
        
        # Recent drafts
        cursor.execute('''
            SELECT username, case_type, created_at 
            FROM drafts 
            ORDER BY created_at DESC 
            LIMIT 10
        ''')
        recent_drafts = cursor.fetchall()
        
        return {
            'total_drafts': total_drafts,
            'drafts_by_type': drafts_by_type,
            'recent_drafts': recent_drafts
        }
    except sqlite3.Error:
        return {'total_drafts': 0, 'drafts_by_type': {}, 'recent_drafts': []}
    finally:
        conn.close()

def save_training_doc(document_name, document_type, content, uploaded_by):
    """Save training document to database"""
    conn = get_db_connection()
    if conn is None:
        return False
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO training_docs (document_name, document_type, content, uploaded_by)
            VALUES (?, ?, ?, ?)
        ''', (document_name, document_type, content, uploaded_by))
        
        conn.commit()
        
        # Log the action
        log_audit_action(st.session_state.user_id, st.session_state.username, "UPLOAD_TRAINING_DOC", f"Uploaded: {document_name}")
        
        return True
    except Exception as e:
        st.error(f"Error saving training doc: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()

def get_training_docs():
    """Get all training documents"""
    conn = get_db_connection()
    if conn is None:
        return []
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT document_name, document_type, uploaded_by, uploaded_at 
            FROM training_docs 
            ORDER BY uploaded_at DESC
        ''')
        
        return cursor.fetchall()
    except sqlite3.Error:
        return []
    finally:
        conn.close()

def get_user_role(username):
    """Get user role"""
    user = get_user_by_username(username)
    return user['role'] if user else "staff"

def check_authentication():
    """Check if user is authenticated"""
    return st.session_state.get('authenticated', False)

def test_database_connection():
    """Test database connection and basic operations with detailed info"""
    try:
        conn = get_db_connection()
        if conn:
            cursor = conn.cursor()
            
            # Get table info
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
            tables = cursor.fetchall()
            
            # Get user count
            cursor.execute("SELECT COUNT(*) FROM users WHERE active = 1")
            active_users = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM users WHERE active = 0")
            inactive_users = cursor.fetchone()[0]
            
            # Get draft count
            cursor.execute("SELECT COUNT(*) FROM drafts")
            draft_count = cursor.fetchone()[0]
            
            conn.close()
            
            return True, f"Database connected. Tables: {[table['name'] for table in tables]}. Active users: {active_users}, Inactive users: {inactive_users}, Drafts: {draft_count}"
        else:
            return False, "Failed to connect to database"
    except Exception as e:
        return False, f"Database test failed: {str(e)}"

def debug_get_all_users_including_inactive():
    """Get all users including inactive ones for debugging"""
    conn = get_db_connection()
    if conn is None:
        return []
    
    try:
        cursor = conn.cursor()
        cursor.execute('''
            SELECT id, username, name, email, role, active, created_at
            FROM users 
            ORDER BY username
        ''')
        
        users = cursor.fetchall()
        return users
    except sqlite3.Error as e:
        print(f"Debug error: {e}")
        return []
    finally:
        conn.close()

# =============================================================================
# DOCUMENT PROCESSOR MODULE
# =============================================================================

class DocumentProcessor:
    @staticmethod
    def extract_text_from_pdf(file):
        try:
            # Reset file pointer to beginning
            file.seek(0)
            
            # Method 1: Try PyMuPDF (fitz)
            try:
                pdf_document = fitz.open(stream=file.read(), filetype="pdf")
                text = ""
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    text += page.get_text()
                pdf_document.close()
                if text.strip():
                    return text
            except Exception as e:
                print(f"PyMuPDF failed: {e}")
            
            # Method 2: Try PyPDF2 as fallback
            file.seek(0)
            try:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                if text.strip():
                    return text
            except Exception as e:
                print(f"PyPDF2 failed: {e}")
            
            return "Unable to extract text from PDF (may be scanned image or encrypted)"
            
        except Exception as e:
            return f"Error extracting PDF text: {str(e)}"
    
    @staticmethod
    def extract_text_from_docx(file):
        try:
            file.seek(0)
            doc = docx.Document(io.BytesIO(file.read()))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        except Exception as e:
            return f"Error extracting DOCX text: {str(e)}"
    
    @staticmethod
    def extract_text_from_txt(file):
        try:
            file.seek(0)
            return file.read().decode('utf-8')
        except Exception as e:
            return f"Error extracting text: {str(e)}"
    
    @staticmethod
    def process_uploaded_file(file):
        file_extension = file.name.split('.')[-1].lower()
        
        # Reset file pointer before processing
        file.seek(0)
        
        if file_extension == 'pdf':
            return DocumentProcessor.extract_text_from_pdf(file)
        elif file_extension == 'docx':
            return DocumentProcessor.extract_text_from_docx(file)
        elif file_extension == 'txt':
            return DocumentProcessor.extract_text_from_txt(file)
        else:
            return f"Unsupported file type: {file_extension}"
    
    @staticmethod
    def chunk_text(text, chunk_size=1000, overlap=200):
        if not text or not text.strip():
            return []
            
        words = text.split()
        chunks = []
        
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            chunks.append(chunk)
            if i + chunk_size >= len(words):
                break
        
        return chunks

# =============================================================================
# GROQ UTILS MODULE
# =============================================================================

class GroqManager:
    """Handles Groq AI draft generation strictly from provided input."""

    def __init__(self):
        # Use environment variable or default key (replace with your actual key)
        self.api_key = os.getenv("GROQ_API_KEY", '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')
        self.client = groq.Groq(api_key=self.api_key)
        self.model_name = "llama-3.3-70b-versatile"

    def generate_draft(self, case_details, medical_summary, retrieved_templates):
        """
        Generate a professional legal demand draft using only user-provided inputs.
        Templates are referenced solely for tone/structure.
        """
        prompt = self._build_prompt(case_details, medical_summary, retrieved_templates)
        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                temperature=0.6,
                top_p=0.95,
                max_tokens=4000,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a professional legal AI assistant that drafts "
                            "high-quality demand letters. "
                            "STRICT RULE: Do not add any information beyond what the user provides. "
                            "Do not fetch external facts. Use templates only for tone and formatting."
                        ),
                    },
                    {"role": "user", "content": prompt},
                ],
            )
            return response.choices[0].message.content.strip()
        except Exception as e:
            return f"Error generating draft: {e}"

    def _build_prompt(self, case_details, medical_summary, retrieved_templates):
        """Construct a strict prompt for AI to generate a factually accurate draft."""
        current_date = datetime.now().strftime("%B %d, %Y")

        plaintiff = case_details.get("plaintiff", "Client")
        defendant = case_details.get("defendant", "Defendant")
        case_type = case_details.get("case_type", "Personal Injury")
        jurisdiction = case_details.get("jurisdiction", "California")
        injury_type = case_details.get("injury_type", "Various injuries")
        treatment_cost = case_details.get("treatment_cost", 0)
        attorney_name = case_details.get("attorney_name", "Legal Counsel")
        law_firm_name = case_details.get("law_firm_name", "Law Office")
        additional_details = case_details.get("additional_details", "")

        template_text = "\n\n--- TEMPLATE REFERENCE ---\n\n".join(retrieved_templates[:3]) if retrieved_templates else "No templates available."

        return f"""
You are a highly specialized legal AI. Generate a **full, professional demand letter** strictly based on the following inputs. 

### ⚠️ STRICT RULES
- ONLY use information provided in the inputs below. Do not invent names, dates, addresses, or details.
- Templates may be used **only for tone and formatting**, not content.
- Ensure the letter is complete, structured, and professional.
- Include a table for damages if applicable.
- Maintain neutral, factual, and persuasive tone.
- Use professional legal formatting: bold section headings, proper tables, clear paragraphs.
- Do NOT use Markdown headings (#, ##, ###).
- Write fully professional, persuasive, and factual demand letter.
- Only use the facts provided; templates are reference for tone/structure only.
- Do not add any external information or assumptions.
- Ensure the letter is self-contained and does not reference external documents.
- Use short paragraphs (4-5 lines each) for readability.
- Don't mention introduction, conclusion, or any other sections explicitly.
- Ensure all tables are cleanly formatted with category and amount columns.
- The letter should be ready to send with no further edits needed.
- If any required information is missing in the template retrieval, **ignore the missing parameters** and generate the draft **fully complete** based on the facts provided.
- Do not use placeholders for missing template parameters; generate the draft naturally with the available data.
- If any parameter (like lost wages or total damages) is missing in the inputs, do NOT write placeholders like "To be determined". 
- Only include rows in the damages table for which values are provided. 
- The draft should still be complete and professional even if some rows are missing.
- If treatment cost is not provided, omit the treatment cost row from the damages table.
- Act as a senior legal editor ensuring clarity, tone, and professionalism.
- Do not add any new facts or assumptions in the rewrite.
- The final output should be a polished, professional demand letter ready for client review.
- if Lost Wages or any other damages are not provided, consider treatment cost as total damages.

### CASE INFORMATION
- Date: {current_date}
- Plaintiff: {plaintiff}
- Defendant: {defendant}
- Jurisdiction: {jurisdiction}
- Case Type: {case_type}
- Injury Type: {injury_type}
- Treatment Cost: ${treatment_cost:,}
- Additional Details: {additional_details}
- Attorney: {attorney_name}
- Law Firm: {law_firm_name}

### MEDICAL SUMMARY
{medical_summary.strip()}

### TEMPLATE STYLE REFERENCE
{template_text}

### OUTPUT REQUIREMENTS
- Bold heading: 'Demand Letter – {case_type} Case'
- Address block: To {defendant}, Jurisdiction: {jurisdiction}
- RE line: {case_type} Claim – {plaintiff} vs. {defendant}
- Sections: Introduction, Statement of Facts, Injuries & Treatment, Damages, Compensation Table, Conclusion
- Signature block with attorney and law firm
- Tables formatted cleanly with category and amount
- Use short paragraphs (3–4 lines each)
- No external facts or assumptions
- Fully self-contained letter
"""

    def rewrite_draft(self, original_draft, improvement_goal):
        """
        Rewrite and improve the draft while strictly keeping facts intact.
        """
        try:
            response = self.client.chat.completions.create(
                model=self.model_name,
                temperature=0.5,
                top_p=0.9,
                max_tokens=3500,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a senior legal editor. Rewrite the provided legal draft to improve it based on the user's goal. "
                            "STRICT RULES:\n"
                            "1. Keep ALL factual information exactly the same\n"
                            "2. Do not add any new facts, numbers, or claims\n"
                            "3. Do not change the meaning or legal arguments\n"
                            "4. Only improve: clarity, tone, professionalism, persuasiveness\n"
                            "5. Maintain the same structure and sections\n"
                            "6. Keep all tables and financial amounts identical\n"
                            "7. Return a complete, polished draft ready for use\n"
                            "8. If you cannot improve it due to constraints, return the original draft unchanged"
                        ),
                    },
                    {
                        "role": "user",
                        "content": f"IMPROVEMENT GOAL: {improvement_goal}\n\n"
                                  f"ORIGINAL DRAFT:\n{original_draft}\n\n"
                                  f"IMPROVED DRAFT:"
                    },
                ],
            )
            result = response.choices[0].message.content.strip()
            
            # Validate that the response is not an error message
            if not result or "error" in result.lower() or "sorry" in result.lower():
                return original_draft  # Return original if AI response is problematic
                
            return result
            
        except Exception as e:
            # Instead of returning an error string, return the original draft
            print(f"AI Rewrite error: {e}")  # For debugging
            return original_draft  # Fallback to original draft

# =============================================================================
# PINECONE UTILS MODULE
# =============================================================================

class PineconeManager:
    def __init__(self):
        self.api_key = "pcsk_34ujt3_Sg56hdoBmpAPJrxubCteDeCrrwNKAhK8aZeHwiWaYkNK9UCpiCtgFR5weBxp9hU"
        self.index_name = "legal-draft-assistant"
        self.model = SentenceTransformer('all-MiniLM-L6-v2')
        self.init_pinecone()
    
    def init_pinecone(self):
        try:
            # Initialize Pinecone with the new SDK
            self.pc = pinecone.Pinecone(api_key=self.api_key)
            
            # Create index if not exists - using FREE TIER supported region
            if self.index_name not in [index.name for index in self.pc.list_indexes()]:
                self.pc.create_index(
                    name=self.index_name,
                    dimension=384,
                    metric="cosine",
                    spec=pinecone.ServerlessSpec(
                        cloud="aws",
                        region="us-east-1"
                    )
                )
            
            # Wait for index to be ready
            import time
            while not self.pc.describe_index(self.index_name).status.ready:
                time.sleep(1)
            
            self.index = self.pc.Index(self.index_name)
            
        except Exception as e:
            raise Exception(f"Pinecone initialization failed: {str(e)}")
    
    def generate_embedding(self, text):
        return self.model.encode(text).tolist()
    
    def store_document(self, text, metadata, document_id=None):
        if document_id is None:
            document_id = str(uuid.uuid4())
        
        embedding = self.generate_embedding(text)
        
        # Prepare metadata
        full_metadata = {
            "text": text[:10000],  # Limit text size for free tier
            **metadata
        }
        
        try:
            # Upsert to Pinecone with error handling
            self.index.upsert(vectors=[(document_id, embedding, full_metadata)])
            return document_id
        except Exception as e:
            print(f"Error storing document: {e}")
            return None
    
    def search_similar(self, query, filter_dict=None, top_k=5):
        try:
            query_embedding = self.generate_embedding(query)
            
            search_params = {
                "vector": query_embedding,
                "top_k": top_k,
                "include_metadata": True
            }
            
            if filter_dict:
                search_params["filter"] = filter_dict
            
            results = self.index.query(**search_params)
            return results
        except Exception as e:
            print(f"Search error: {e}")
            return type('obj', (object,), {'matches': []})()
    
    def delete_document(self, document_id):
        try:
            self.index.delete(ids=[document_id])
        except Exception as e:
            print(f"Delete error: {e}")
    
    def get_index_stats(self):
        """Get index statistics"""
        try:
            return self.index.describe_index_stats()
        except Exception as e:
            print(f"Stats error: {e}")
            return {}

# =============================================================================
# MAIN APPLICATION
# =============================================================================

# Manager factories (session-safe)
def get_pinecone_manager():
    if 'pinecone_mgr' not in st.session_state:
        try:
            st.session_state.pinecone_mgr = PineconeManager()
        except Exception as e:
            st.error(f"Failed to initialize Pinecone: {e}")
            # Create a mock manager that won't crash the app
            st.session_state.pinecone_mgr = type('obj', (object,), {
                'search_similar': lambda *args, **kwargs: type('obj', (object,), {'matches': []})(),
                'store_document': lambda *args, **kwargs: None,
                'get_index_stats': lambda: {}
            })()
    return st.session_state.pinecone_mgr

def get_ai_manager():
    if 'ai_mgr' not in st.session_state:
        try:
            st.session_state.ai_mgr = GroqManager()
        except Exception as e:
            st.error(f"Failed to initialize AI: {e}")
            # Create a mock manager
            st.session_state.ai_mgr = type('obj', (object,), {
                'generate_draft': lambda *args, **kwargs: "AI service unavailable. Please check your API keys.",
                'rewrite_draft': lambda original, goal: original
            })()
    return st.session_state.ai_mgr

def get_document_processor():
    if 'doc_processor' not in st.session_state:
        st.session_state.doc_processor = DocumentProcessor()
    return st.session_state.doc_processor

# Utility functions
def check_ai_setup():
    """Check if Groq is properly configured"""
    try:
        ai_mgr = get_ai_manager()
        # Simple test query
        _ = ai_mgr.client.chat.completions.create(
            messages=[{"role": "user", "content": "ping"}],
            model=ai_mgr.model_name,
            max_tokens=5
        )
        return True, "✅ Groq API is working"
    except Exception as e:
        return False, f"❌ Groq setup failed: {str(e)}"

def format_draft_with_tables(draft_text):
    """Convert markdown tables to properly formatted text tables"""
    draft_text = re.sub(r'^\|[-:\s|]+\|\s*$', '', draft_text, flags=re.MULTILINE)
    table_pattern = r'((?:\|.*\|\n)+)'

    def replace_table(match):
        table_text = match.group(0)
        lines = table_text.strip().split('\n')
        formatted_lines = []
        for line in lines:
            if re.match(r'^\|[-:\s|]+\|\s*$', line.strip()):
                continue
            if '|' in line:
                cells = [cell.strip() for cell in line.strip('|').split('|')]
                formatted_lines.append(" | ".join(cells))
        if formatted_lines:
            return "\n" + "\n".join(formatted_lines) + "\n"
        return table_text

    formatted_draft = re.sub(table_pattern, replace_table, draft_text, flags=re.MULTILINE)
    return formatted_draft

def rewrite_draft_with_ai(original_draft, improvement_goal):
    """Use AI to rewrite and improve the draft with better error handling"""
    ai_mgr = get_ai_manager()
    try:
        improved_draft = ai_mgr.rewrite_draft(original_draft, improvement_goal)
        
        # Check if we got a valid draft back (not empty and not an error message)
        if (improved_draft and 
            improved_draft.strip() and 
            not improved_draft.startswith("Error") and
            "error generating" not in improved_draft.lower() and
            "error rewriting" not in improved_draft.lower()):
            
            improved_draft = format_draft_with_tables(improved_draft)
            return improved_draft
        else:
            # If AI returned an error or empty response, return original
            return original_draft
            
    except Exception as e:
        print(f"Rewrite error: {e}")  # For debugging
        return original_draft  # Return original draft on any error

def create_fallback_draft(case_details, medical_text):
    """Create a fallback draft when AI unavailable"""
    current_date = datetime.now().strftime("%B %d, %Y")
    medical_preview = medical_text[:500] + "..." if len(medical_text) > 500 else medical_text
    return f"""
**DEMAND LETTER - {case_details.get('case_type', 'Personal Injury').upper()} CASE**

{current_date}

TO:
{case_details.get('defendant', 'Defendant')}
Jurisdiction: {case_details.get('jurisdiction', 'California')}

RE: {case_details.get('case_type', 'Personal Injury')} Claim - {case_details.get('plaintiff', 'Plaintiff')} vs. {case_details.get('defendant', 'Defendant')}

Dear Claims Adjuster,

This law firm represents {case_details.get('plaintiff', 'our client')} in connection with injuries sustained due to your negligence.

**INCIDENT SUMMARY**
{case_details.get('additional_details', 'Our client suffered injuries as a result of your negligent conduct.')}

**INJURIES AND MEDICAL TREATMENT**
Our client sustained {case_details.get('injury_type', 'serious injuries')} requiring medical attention.

Medical Summary:
{medical_preview}

The total medical expenses incurred amount to ${case_details.get('treatment_cost', '10,000'):,}.

**DAMAGES BREAKDOWN**
| Category | Amount (USD) |
|----------|-------------:|
| Medical Expenses | ${case_details.get('treatment_cost', '10,000'):,} |
| Pain and Suffering | $15,000 |
| Lost Wages | $5,000 |
| **TOTAL DEMAND** | **$30,000** |

**DEMAND FOR COMPENSATION**
We hereby demand payment of $30,000 to fully compensate our client for medical expenses, pain and suffering, and other damages resulting from this incident.

Please forward this demand to your insurance carrier and have them contact us within 30 days to resolve this matter. Failure to respond may compel us to pursue legal action.

Sincerely,

{case_details.get('attorney_name', 'Legal Counsel')}
{case_details.get('law_firm_name', 'Law Offices of Legal Counsel')}
"""

def create_professional_word_download(text, filename):
    """Create a professional Word document with proper formatting like MS Word"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import io
        
        def set_cell_background(cell, color):
            """Set cell background color"""
            tcPr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), color)
            tcPr.append(shd)
        
        def set_cell_border(cell):
            """Set cell borders"""
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            # Add all borders
            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)
        
        def clean_unicode_text(text):
            """Replace Unicode characters with ASCII equivalents"""
            replacements = {
                '\u2013': '-',  # en dash
                '\u2014': '-',  # em dash
                '\u2018': "'",  # left single quote
                '\u2019': "'",  # right single quote
                '\u201C': '"',  # left double quote
                '\u201D': '"',  # right double quote
                '\u2022': '•',  # bullet (keep this one)
                '\u2026': '...', # ellipsis
            }
            
            cleaned_text = str(text)
            for unicode_char, ascii_replacement in replacements.items():
                cleaned_text = cleaned_text.replace(unicode_char, ascii_replacement)
            
            return cleaned_text
        
        doc = Document()
        
        # Set document margins (like Word default)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title - Centered and bold
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("LEGAL DEMAND DRAFT")
        title_run.font.size = Pt(16)
        title_run.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
        title_run.bold = True
        title_run.font.name = 'Arial'
        
        doc.add_paragraph()  # Add space
        
        # Preprocess text
        processed_text = clean_unicode_text(text)
        lines = processed_text.split('\n')
        current_table = []
        in_table = False
        
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()  # Add empty paragraph for blank lines
                continue
                
            # Handle headings (bold text with **)
            if line.startswith('**') and line.endswith('**'):
                heading_text = line.strip('*').strip()
                heading = doc.add_paragraph()
                heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
                heading_run = heading.add_run(heading_text)
                heading_run.font.size = Pt(14)
                heading_run.font.color.rgb = RGBColor(0, 0, 128)  # Navy blue
                heading_run.bold = True
                heading_run.font.name = 'Arial'
                heading.paragraph_format.space_after = Pt(6)
                continue
                
            # Handle table rows
            if '|' in line and any(cell.strip() for cell in line.split('|')):
                if not in_table:
                    current_table = []
                    in_table = True
                
                cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                if cells:
                    current_table.append(cells)
                continue
            else:
                # Process accumulated table
                if in_table and current_table and len(current_table) > 1:
                    max_cols = max(len(row) for row in current_table)
                    table = doc.add_table(rows=len(current_table), cols=max_cols)
                    table.style = 'Table Grid'
                    
                    # Set column widths
                    for col in table.columns:
                        col.width = Inches(2.5)
                    
                    for i, row in enumerate(current_table):
                        for j in range(max_cols):
                            cell_text = row[j] if j < len(row) else ""
                            cell = table.cell(i, j)
                            cell.text = cell_text
                            
                            # Set borders for all cells
                            set_cell_border(cell)
                            
                            # Style header row
                            if i == 0:
                                set_cell_background(cell, "2F75B5")  # Blue header
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in paragraph.runs:
                                        run.bold = True
                                        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(11)
                            else:
                                # Alternate row colors for data rows
                                if i % 2 == 1:  # Even rows
                                    set_cell_background(cell, "DDEBF7")  # Light blue
                                else:  # Odd rows
                                    set_cell_background(cell, "FFFFFF")  # White
                                
                                for paragraph in cell.paragraphs:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in paragraph.runs:
                                        run.font.name = 'Arial'
                                        run.font.size = Pt(10)
                                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
                    
                    doc.add_paragraph()  # Add space after table
                    in_table = False
                    current_table = []
                
                # Regular paragraph
                if line.strip():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(line)
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
        
        # Handle case where text ends with a table
        if in_table and current_table and len(current_table) > 1:
            max_cols = max(len(row) for row in current_table)
            table = doc.add_table(rows=len(current_table), cols=max_cols)
            table.style = 'Table Grid'
            
            for col in table.columns:
                col.width = Inches(2.5)
            
            for i, row in enumerate(current_table):
                for j in range(max_cols):
                    cell_text = row[j] if j < len(row) else ""
                    cell = table.cell(i, j)
                    cell.text = cell_text
                    set_cell_border(cell)
                    
                    if i == 0:
                        set_cell_background(cell, "2F75B5")
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.name = 'Arial'
                                run.font.size = Pt(11)
                    else:
                        if i % 2 == 1:
                            set_cell_background(cell, "DDEBF7")
                        else:
                            set_cell_background(cell, "FFFFFF")
                        
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(0, 0, 0)
        
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    except Exception as e:
        st.error(f"Word document generation failed: {e}")
        # Fallback: return text as plain file
        buffer = io.BytesIO()
        buffer.write(text.encode('utf-8'))
        buffer.seek(0)
        return buffer

def create_professional_pdf_download(text, filename):
    """Create a professional PDF document that handles Unicode characters properly"""
    try:
        from fpdf import FPDF
        import io
        
        class LegalPDF(FPDF):
            def __init__(self):
                super().__init__()
                self.set_auto_page_break(auto=True, margin=15)
                self.set_margins(left=20, top=20, right=20)
            
            def header(self):
                # Title - Centered and colored
                self.set_fill_color(47, 117, 181)  # Blue background
                self.set_text_color(255, 255, 255)  # White text
                self.set_font('Arial', 'B', 16)
                self.cell(0, 12, 'LEGAL DEMAND DRAFT', 0, 1, 'C', fill=True)
                self.ln(10)
                self.set_text_color(0, 0, 0)  # Black text
            
            def footer(self):
                self.set_y(-15)
                self.set_font('Arial', 'I', 9)
                self.set_text_color(128, 128, 128)
                self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')
            
            def add_heading(self, text, level=1):
                self.set_font('Arial', 'B', 14 if level == 1 else 12)
                self.set_text_color(0, 0, 128)  # Navy blue
                clean_text = self._clean_text(text)
                self.multi_cell(0, 8, clean_text)
                self.ln(4)
                self.set_text_color(0, 0, 0)
            
            def add_paragraph(self, text):
                self.set_font('Arial', '', 11)
                if text and text.strip():
                    clean_text = self._clean_text(text)
                    # Justify text like Word
                    self.multi_cell(0, 6, clean_text)
                    self.ln(3)
            
            def add_table(self, rows):
                if not rows or len(rows) < 2:
                    return
                    
                col_count = max(len(row) for row in rows)
                page_width = self.w - 40  # Account for margins
                col_width = page_width / col_count
                
                # Header row with blue background
                self.set_fill_color(47, 117, 181)  # Blue
                self.set_text_color(255, 255, 255)  # White
                self.set_font('Arial', 'B', 11)
                
                for header in rows[0]:
                    clean_header = self._clean_text(str(header))
                    self.cell(col_width, 10, clean_header, border=1, align='C', fill=True)
                self.ln(10)
                
                # Data rows with alternating colors
                self.set_font('Arial', '', 10)
                self.set_text_color(0, 0, 0)  # Black text
                
                for i, row in enumerate(rows[1:]):
                    # Alternate row colors
                    if i % 2 == 0:
                        self.set_fill_color(221, 235, 247)  # Light blue
                    else:
                        self.set_fill_color(255, 255, 255)  # White
                    
                    for j in range(col_count):
                        cell_text = row[j] if j < len(row) else ""
                        clean_cell = self._clean_text(str(cell_text))
                        if len(clean_cell) > 30:
                            clean_cell = clean_cell[:27] + "..."
                        self.cell(col_width, 8, clean_cell, border=1, align='C', fill=True)
                    
                    self.ln(8)
                
                self.ln(5)
            
            def _clean_text(self, text):
                """Replace Unicode characters with ASCII equivalents for PDF compatibility"""
                if not text:
                    return ""
                
                # Common Unicode character replacements
                replacements = {
                    '\u2013': '-',  # en dash
                    '\u2014': '-',  # em dash
                    '\u2018': "'",  # left single quote
                    '\u2019': "'",  # right single quote
                    '\u201C': '"',  # left double quote
                    '\u201D': '"',  # right double quote
                    '\u2022': '*',  # bullet
                    '\u2026': '...', # ellipsis
                    '\u20AC': 'EUR', # euro sign
                    '\u00A3': 'GBP', # pound sign
                    '\u00A9': '(c)', # copyright
                    '\u00AE': '(R)', # registered
                    '\u2122': '(TM)', # trademark
                }
                
                cleaned_text = str(text)
                for unicode_char, ascii_replacement in replacements.items():
                    cleaned_text = cleaned_text.replace(unicode_char, ascii_replacement)
                
                return cleaned_text
        
        def preprocess_text_for_pdf(text):
            """Preprocess the entire text to handle Unicode characters"""
            replacements = {
                '\u2013': '-',
                '\u2014': '-',
                '\u2018': "'",
                '\u2019': "'",
                '\u201C': '"',
                '\u201D': '"',
                '\u2022': '*',
                '\u2026': '...',
                '\u20AC': 'EUR',
                '\u00A3': 'GBP',
                '\u00A9': '(c)',
                '\u00AE': '(R)',
                '\u2122': '(TM)',
            }
            
            cleaned_text = str(text)
            for unicode_char, ascii_replacement in replacements.items():
                cleaned_text = cleaned_text.replace(unicode_char, ascii_replacement)
            
            return cleaned_text
        
        # Always return a valid buffer
        try:
            # Preprocess the entire text first
            processed_text = preprocess_text_for_pdf(text)
            
            pdf = LegalPDF()
            pdf.add_page()
            
            if not processed_text or not processed_text.strip():
                pdf.add_paragraph("No content available for this draft.")
            else:
                lines = processed_text.split('\n')
                current_table = []
                in_table = False
                
                for line in lines:
                    line = line.strip()
                    if not line:
                        pdf.ln(3)  # Add space for blank lines
                        continue
                    
                    # Handle headings (bold text with **)
                    if line.startswith('**') and line.endswith('**'):
                        heading_text = line.strip('*').strip()
                        if heading_text:
                            pdf.add_heading(heading_text, level=1)
                        continue
                    
                    # Handle table rows
                    if '|' in line and any(cell.strip() for cell in line.split('|')):
                        if not in_table:
                            current_table = []
                            in_table = True
                        
                        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
                        if cells:
                            current_table.append(cells)
                        continue
                    else:
                        # Process table
                        if in_table and current_table and len(current_table) > 1:
                            pdf.add_table(current_table)
                            in_table = False
                            current_table = []
                        
                        # Regular paragraph
                        if (line.strip() and 
                            not re.match(r'^[\s\|:-]+$', line.strip()) and
                            not line.strip().startswith('|---')):
                            pdf.add_paragraph(line)
                
                # Final table
                if in_table and current_table and len(current_table) > 1:
                    pdf.add_table(current_table)
            
            # Generate PDF bytes
            pdf_bytes = pdf.output(dest='S').encode('latin-1', 'replace')
            buffer = io.BytesIO(pdf_bytes)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            st.error(f"PDF generation error: {str(e)}")
            # Ultimate fallback - simple PDF
            try:
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font('Arial', 'B', 16)
                pdf.cell(0, 10, 'LEGAL DEMAND DRAFT', 0, 1, 'C')
                pdf.ln(10)
                pdf.set_font('Arial', '', 12)
                
                safe_text = preprocess_text_for_pdf(text) if text else "No content available"
                pdf.multi_cell(0, 8, safe_text[:1500])
                
                pdf_bytes = pdf.output(dest='S').encode('latin-1', 'replace')
                buffer = io.BytesIO(pdf_bytes)
                buffer.seek(0)
                return buffer
            except Exception as final_error:
                # Last resort - empty but valid PDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font('Arial', 'B', 16)
                pdf.cell(0, 10, 'LEGAL DEMAND DRAFT', 0, 1, 'C')
                pdf_bytes = pdf.output(dest='S').encode('latin-1', 'replace')
                buffer = io.BytesIO(pdf_bytes)
                buffer.seek(0)
                return buffer
    except Exception as e:
        st.error(f"PDF generation failed: {e}")
        # Fallback: return text as plain file
        buffer = io.BytesIO()
        buffer.write(text.encode('utf-8'))
        buffer.seek(0)
        return buffer

# Initialize session state
def initialize_session_state():
    """Initialize all session state variables"""
    if 'draft_history' not in st.session_state:
        st.session_state.draft_history = []
    if 'current_draft' not in st.session_state:
        st.session_state.current_draft = ""
    if 'medical_text' not in st.session_state:
        st.session_state.medical_text = ""
    if 'last_snapshot' not in st.session_state:
        st.session_state.last_snapshot = {}
    if 'rewrite_goals' not in st.session_state:
        st.session_state.rewrite_goals = [
            "Make it more persuasive and compelling",
            "Improve legal language and professionalism",
            "Enhance clarity and readability",
            "Strengthen the demand justification",
            "Make it more concise and direct"
        ]
    if 'ai_rewrite_trigger' not in st.session_state:
        st.session_state.ai_rewrite_trigger = False
    if 'ai_rewrite_goal' not in st.session_state:
        st.session_state.ai_rewrite_goal = ""

# Training functions
def process_training_files(files, document_type):
    pinecone_mgr = get_pinecone_manager()
    doc_processor = get_document_processor()
    
    progress_bar = st.progress(0)
    status_text = st.empty()
    successful_uploads = 0
    
    for i, file in enumerate(files):
        status_text.text(f"Processing {file.name}...")
        
        try:
            text = doc_processor.process_uploaded_file(file)
            
            if text and "Error" not in text and "Unable" not in text:
                # Store in Pinecone
                metadata = {
                    "document_type": document_type,
                    "file_name": file.name,
                    "upload_date": datetime.now().isoformat(),
                    "uploaded_by": st.session_state.username,
                    "category": "legal_draft"
                }
                
                result = pinecone_mgr.store_document(text, metadata)
                
                # Also save to SQL database
                db_success = save_training_doc(file.name, document_type, text, st.session_state.username)
                
                if result and db_success:
                    successful_uploads += 1
                    st.success(f"✅ Successfully stored {file.name} as {document_type}")
                else:
                    st.error(f"❌ Failed to store {file.name}")
            else:
                st.error(f"❌ Failed to extract text from {file.name}: {text}")
        
        except Exception as e:
            st.error(f"❌ Error processing {file.name}: {str(e)}")
        
        progress_bar.progress((i + 1) / len(files))
    
    status_text.text(f"Processing complete! {successful_uploads}/{len(files)} files stored.")
    
    if successful_uploads > 0:
        st.success(f"🎉 Successfully stored {successful_uploads} {document_type} files!")

def show_vector_db_status():
    pinecone_mgr = get_pinecone_manager()
    
    try:
        stats = pinecone_mgr.get_index_stats()
        
        st.write("### Vector Database Status")
        st.write(f"Total Vectors: {stats.get('total_vector_count', 0)}")
        st.write(f"Index Dimension: {stats.get('dimension', 0)}")
        
        st.write("### All Documents in Database")
        all_results = pinecone_mgr.search_similar("legal medical demand", top_k=20)
        
        if not all_results.matches:
            st.error("❌ NO DOCUMENTS FOUND IN DATABASE!")
            st.info("Please upload templates in the 'Train Model' tab")
            return
            
        for i, match in enumerate(all_results.matches):
            doc_type = match.metadata.get('document_type', 'UNKNOWN')
            file_name = match.metadata.get('file_name', 'UNKNOWN')
            with st.expander(f"{doc_type.upper()}: {file_name} (Score: {match.score:.3f})"):
                st.write(f"**Type:** {doc_type}")
                st.write(f"**File:** {file_name}")
                st.text(match.metadata.get('text', '')[:300] + "..." if len(match.metadata.get('text', '')) > 300 else match.metadata.get('text', ''))
    
    except Exception as e:
        st.error(f"Error accessing vector database: {e}")

# User Management Functions
def show_user_management():
    """User management interface for admin with proper database integration"""
    st.subheader("User Management")
    
    # Test database connection
    if st.button("🔧 Test Database Connection"):
        success, message = test_database_connection()
        if success:
            st.success(message)
        else:
            st.error(message)
    
    # Create new user
    with st.expander("➕ Create New User", expanded=True):
        with st.form("create_user_form", clear_on_submit=True):
            col1, col2 = st.columns(2)
            
            with col1:
                new_username = st.text_input("Username*", placeholder="Enter username")
                new_password = st.text_input("Password*", type="password", placeholder="Enter password")
                confirm_password = st.text_input("Confirm Password*", type="password", placeholder="Confirm password")
            
            with col2:
                new_name = st.text_input("Full Name*", placeholder="Enter full name")
                new_email = st.text_input("Email*", placeholder="Enter email")
                new_role = st.selectbox("Role*", ["staff", "admin"])
            
            create_button = st.form_submit_button("Create User")
            
            if create_button:
                if not all([new_username, new_password, confirm_password, new_name, new_email]):
                    st.error("Please fill all required fields (*)")
                elif new_password != confirm_password:
                    st.error("Passwords do not match")
                elif len(new_password) < 6:
                    st.error("Password must be at least 6 characters")
                else:
                    success, message = create_user(
                        username=new_username,
                        password=new_password,
                        name=new_name,
                        email=new_email,
                        role=new_role,
                        created_by=st.session_state.user_id
                    )
                    if success:
                        st.success(message)
                        st.rerun()
                    else:
                        st.error(message)
    
    # User list and management
    st.subheader("Existing Users")
    
    # Refresh users list
    users = get_all_users()
    
    if not users:
        st.info("No users found in database")
        return
    
    # Display users in a table with actions
    user_data = []
    for username, user_info in users.items():
        user_data.append({
            "ID": user_info['id'],
            "Username": username,
            "Name": user_info.get('name', ''),
            "Email": user_info.get('email', ''),
            "Role": user_info.get('role', 'staff'),
            "Created": user_info.get('created_at', ''),
            "Created By": user_info.get('created_by', 'system'),
            "Last Login": user_info.get('last_login', 'Never')
        })
    
    if user_data:
        # Display user count
        st.write(f"**Total Active Users:** {len(user_data)}")
        
        # User table
        df = pd.DataFrame(user_data)
        st.dataframe(df, use_container_width=True, hide_index=True)
        
        # User actions in separate expanders for better organization
        st.subheader("User Actions")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            with st.expander("✏️ Edit User", expanded=True):
                edit_users = [user for user in user_data if user["Username"] != "admin"]
                if edit_users:
                    edit_username = st.selectbox("Select user to edit", 
                                               [user["Username"] for user in edit_users],
                                               key="edit_select")
                    if edit_username:
                        user_to_edit = users[edit_username]
                        with st.form(f"edit_form_{edit_username}"):
                            st.write(f"Editing: **{edit_username}**")
                            edit_name = st.text_input("Full Name", value=user_to_edit.get('name', ''), key=f"name_{edit_username}")
                            edit_email = st.text_input("Email", value=user_to_edit.get('email', ''), key=f"email_{edit_username}")
                            edit_role = st.selectbox("Role", ["staff", "admin"], 
                                                   index=0 if user_to_edit.get('role') == "staff" else 1,
                                                   key=f"role_{edit_username}")
                            
                            if st.form_submit_button("💾 Update User", use_container_width=True):
                                success, message = update_user(
                                    user_to_edit['id'],
                                    name=edit_name,
                                    email=edit_email,
                                    role=edit_role
                                )
                                if success:
                                    st.success(message)
                                    st.rerun()
                                else:
                                    st.error(message)
                else:
                    st.info("No users available to edit")
        
        with col2:
            with st.expander("🔑 Change Password", expanded=True):
                pwd_users = [user for user in user_data]
                if pwd_users:
                    pwd_username = st.selectbox("Select user", 
                                              [user["Username"] for user in pwd_users],
                                              key="pwd_select")
                    if pwd_username:
                        user_to_pwd = users[pwd_username]
                        with st.form(f"pwd_form_{pwd_username}"):
                            st.write(f"Changing password for: **{pwd_username}**")
                            new_pwd = st.text_input("New Password", type="password", key=f"new_pwd_{pwd_username}")
                            confirm_pwd = st.text_input("Confirm Password", type="password", key=f"confirm_pwd_{pwd_username}")
                            
                            if st.form_submit_button("🔐 Change Password", use_container_width=True):
                                if not new_pwd or not confirm_pwd:
                                    st.error("Please enter both password fields")
                                elif new_pwd != confirm_pwd:
                                    st.error("Passwords do not match")
                                elif len(new_pwd) < 6:
                                    st.error("Password must be at least 6 characters")
                                else:
                                    success, message = update_user(user_to_pwd['id'], password=new_pwd)
                                    if success:
                                        st.success("Password updated successfully")
                                        st.rerun()
                                    else:
                                        st.error(message)
                else:
                    st.info("No users available")
        
        with col3:
            with st.expander("🗑️ Delete User", expanded=True):
                del_users = [user for user in user_data if user["Username"] != "admin"]
                if del_users:
                    del_username = st.selectbox("Select user to delete", 
                                              [user["Username"] for user in del_users],
                                              key="del_select")
                    if del_username:
                        user_to_del = users[del_username]
                        st.warning(f"⚠️ You are about to delete user: **{del_username}**")
                        st.write(f"Name: {user_to_del.get('name', 'N/A')}")
                        st.write(f"Email: {user_to_del.get('email', 'N/A')}")
                        st.write(f"User ID: {user_to_del['id']}")
                        
                        # Add confirmation with user ID for extra safety
                        confirm_text = st.text_input(
                            f"Type 'DELETE {del_username}' to confirm:",
                            placeholder=f"DELETE {del_username}",
                            key=f"confirm_delete_{del_username}"
                        )
                        
                        if st.button("🚫 Confirm Delete User", type="secondary", use_container_width=True, key=f"delete_btn_{del_username}"):
                            if confirm_text == f"DELETE {del_username}":
                                success, message = delete_user(user_to_del['id'], del_username)
                                if success:
                                    st.success(message)
                                    # Force refresh by rerunning
                                    st.rerun()
                                else:
                                    st.error(message)
                            else:
                                st.error("Confirmation text does not match. Please type exactly as shown.")
                else:
                    st.info("No users available to delete")
    
    # Refresh button
    if st.button("🔄 Refresh User List"):
        st.rerun()

def show_audit_log():
    """Show audit log of user activities"""
    st.subheader("Audit Log - System Activities")
    
    logs = get_audit_logs(limit=50)
    
    if not logs:
        st.info("No audit logs available")
        return
    
    # Convert logs to dataframe
    log_data = []
    for log in logs:
        log_data.append({
            "Timestamp": log[0],
            "Username": log[1],
            "Action": log[2],
            "Details": log[3],
            "User Name": log[4] if log[4] else log[1]
        })
    
    df = pd.DataFrame(log_data)
    st.dataframe(df, use_container_width=True)
    
    # Export audit log
    if st.button("Export Audit Log to CSV"):
        csv = df.to_csv(index=False)
        st.download_button(
            label="Download CSV",
            data=csv,
            file_name=f"audit_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
            mime="text/csv"
        )

def show_training_docs():
    """Show training documents from database"""
    st.subheader("Training Documents")
    
    docs = get_training_docs()
    
    if not docs:
        st.info("No training documents uploaded yet")
        return
    
    doc_data = []
    for doc in docs:
        doc_data.append({
            "Document Name": doc[0],
            "Type": doc[1],
            "Uploaded By": doc[2],
            "Uploaded At": doc[3]
        })
    
    df = pd.DataFrame(doc_data)
    st.dataframe(df, use_container_width=True)
    
    st.subheader("Draft Statistics")
    stats = get_all_drafts_stats()
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Drafts", stats['total_drafts'])
    with col2:
        st.metric("Case Types", len(stats['drafts_by_type']))
    with col3:
        st.metric("Active Users", len(get_all_users()))
    
    # Drafts by case type
    if stats['drafts_by_type']:
        st.write("### Drafts by Case Type")
        type_df = pd.DataFrame(list(stats['drafts_by_type'].items()), columns=['Case Type', 'Count'])
        st.bar_chart(type_df.set_index('Case Type'))

# Admin dashboard
def show_admin_dashboard():
    st.title("👨‍💼 Admin Dashboard")
    st.info("Manage users, upload training data, and monitor system activity.")
    
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["👥 User Management", "📤 Train Model", "📊 Database Status", "📋 Audit Log", "📈 Analytics"])
    
    with tab1:
        show_user_management()
    
    with tab2:
        st.subheader("Upload Templates")
        files = st.file_uploader("Choose template files", type=["txt","docx","pdf"], accept_multiple_files=True, key="admin_upload")
        if files and st.button("Process & Store Templates"):
            process_training_files(files, "template")
        
        # Show training documents
        show_training_docs()
    
    with tab3:
        st.subheader("Vector Database Status")
        if st.button("Check Database Status"):
            show_vector_db_status()
    
    with tab4:
        show_audit_log()
    
    with tab5:
        show_training_docs()

# Staff dashboard
def show_staff_dashboard():
    st.title("👩‍💻 Staff Dashboard")
    tab1, tab2 = st.tabs(["📝 Generate Draft", "📋 Draft History"])

    with tab1:
        show_draft_generation_interface("staff")
    with tab2:
        st.header("Your Draft History")
        user_drafts = [d for d in st.session_state.draft_history if d['user'] == st.session_state.username]
        if not user_drafts:
            st.info("No drafts generated yet")
        else:
            for i, draft in enumerate(reversed(user_drafts)):
                with st.expander(f"Draft {len(user_drafts)-i} - {draft['timestamp']}"):
                    st.write(f"**Case Type:** {draft['case_type']}")
                    st.text_area("Draft Content", draft['content'], height=200, key=f"history_{i}")

# Draft generation interface
def show_draft_generation_interface(user_type):
    st.header("Generate Demand Draft")

    # Case Details - Better aligned form
    with st.container():
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("👤 Party Information")
            plaintiff = st.text_input("**Plaintiff Name**", placeholder="Sarah Davis", key="plaintiff_input")
            defendant = st.text_input("**Defendant Name**", placeholder="Mark Thompson", key="defendant_input")
            case_type = st.selectbox(
                "**Case Type**",
                ["Personal Injury", "Medical Malpractice", "Workers Compensation", "Auto Accident", "Other"],
                key="case_type_input"
            )
        
        with col2:
            st.subheader("📍 Case Details")
            jurisdiction = st.selectbox(
                "**Jurisdiction**",
                ["California", "New York", "Texas", "Florida", "Federal", "Other"],
                key="jurisdiction_input"
            )
            injury_type = st.text_input("**Injury Type**", placeholder="Whiplash, Mild Concussion", key="injury_type_input")
            treatment_cost = st.number_input("**Treatment Cost ($)**", min_value=0, value=10000, key="treatment_cost_input")

    # Additional Details
    st.subheader("📋 Incident Details")
    additional_details = st.text_area(
        "**Describe the incident, timeline, and relevant information:**",
        placeholder="e.g., Client was rear-ended at a red light on June 5, 2024. The impact caused immediate neck pain and headache. Client was transported via ambulance to emergency room...",
        height=100,
        key="additional_details_input"
    )

    # Attorney Information
    st.subheader("⚖️ Attorney & Law Firm")
    col1, col2 = st.columns(2)
    with col1:
        attorney_name = st.text_input("**Attorney Name**", placeholder="Elizabeth Grant, Esq.", key="attorney_input")
    with col2:
        law_firm_name = st.text_input("**Law Firm Name**", placeholder="Grant & Associates Legal Group, LLP", key="lawfirm_input")

    if not attorney_name:
        attorney_name = "Legal Counsel"
    if not law_firm_name:
        law_firm_name = "Law Office of Justice & Associates"

    # Medical Records
    st.subheader("🏥 Medical Records")
    
    medical_files = st.file_uploader(
        "**Upload Medical Records (PDF/DOCX/TXT)**",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        key=f"medical_upload_{user_type}",
        help="Upload medical reports, bills, and treatment records"
    )

    medical_text = st.text_area(
        "**Or enter medical summary directly:**",
        placeholder="Patient presented with neck pain and headache following motor vehicle accident. Emergency room treatment included cervical spine X-rays and pain management. Diagnosed with whiplash and mild concussion. Follow-up treatment included 4 weeks of chiropractic therapy...",
        value=st.session_state.get("medical_text", ""),
        height=150,
        key=f"medical_summary_{user_type}"
    )

    st.session_state.medical_text = medical_text.strip()

    if medical_files:
        processed_text = ""
        for file in medical_files:
            text = get_document_processor().process_uploaded_file(file)
            processed_text += f"\n\n--- {file.name} ---\n{text}"
        st.session_state.medical_text = processed_text
        with st.expander("📄 Extracted Medical Text", expanded=False):
            st.text_area("Medical Records Content", processed_text, height=200, key="extracted_medical")

    # Detect input changes
    current_snapshot = {
        "plaintiff": plaintiff.strip(),
        "defendant": defendant.strip(),
        "case_type": case_type,
        "jurisdiction": jurisdiction,
        "injury_type": injury_type.strip(),
        "treatment_cost": treatment_cost,
        "additional_details": additional_details.strip(),
        "medical_text": st.session_state.medical_text.strip()
    }

    if current_snapshot != st.session_state.last_snapshot:
        st.session_state.current_draft = ""
        st.session_state.last_snapshot = current_snapshot.copy()

    # Generate Draft Button
    st.markdown("---")
    if st.button("🤖 **GENERATE DRAFT**", type="primary", use_container_width=True):
        if not plaintiff or not defendant:
            st.error("❌ Please fill in Plaintiff and Defendant names.")
            return
        if not st.session_state.medical_text.strip():
            st.error("❌ Please provide medical records or a summary before generating.")
            return

        case_details = {
            "plaintiff": plaintiff,
            "defendant": defendant,
            "case_type": case_type,
            "jurisdiction": jurisdiction,
            "injury_type": injury_type,
            "treatment_cost": treatment_cost,
            "additional_details": additional_details,
            "attorney_name": attorney_name,
            "law_firm_name": law_firm_name
        }

        with st.spinner("🔍 Searching templates and generating professional draft..."):
            try:
                pinecone_mgr = get_pinecone_manager()
                ai_mgr = get_ai_manager()
                query_text = f"{case_type.lower()} demand letter template"

                # Retrieve templates
                results = pinecone_mgr.search_similar(
                    query_text,
                    filter_dict={"document_type": {"$in": ["template"]}},
                    top_k=8,
                )

                retrieved_templates = []
                for match in getattr(results, "matches", []):
                    meta = getattr(match, "metadata", {})
                    score = getattr(match, "score", 0)
                    if score > 0.5 and "text" in meta:
                        retrieved_templates.append(meta["text"])

                # Fallback if no templates
                if not retrieved_templates:
                    fallback_results = pinecone_mgr.search_similar(
                        f"{case_type} demand letter {jurisdiction}", top_k=6
                    )
                    for match in getattr(fallback_results, "matches", []):
                        meta = getattr(match, "metadata", {})
                        score = getattr(match, "score", 0)
                        if score > 0.45 and "text" in meta:
                            retrieved_templates.append(meta["text"])

                # Generate draft via AI
                draft = ai_mgr.generate_draft(case_details, st.session_state.medical_text, retrieved_templates)
                if "Error generating draft" in draft:
                    st.warning("⚠️ AI generation failed, using fallback template.")
                    draft = create_fallback_draft(case_details, st.session_state.medical_text)
                else:
                    draft = format_draft_with_tables(draft)

                # Store draft in session state
                st.session_state.current_draft = draft
                
                # Save draft to database
                draft_id = save_draft_to_db(
                    st.session_state.user_id,
                    st.session_state.username,
                    case_type,
                    draft,
                    st.session_state.medical_text,
                    case_details
                )
                
                if draft_id:
                    # Reload drafts from database to ensure consistency
                    load_user_drafts(st.session_state.user_id)
                    st.success("✅ Professional draft generated and saved successfully!")
                else:
                    st.error("❌ Draft generated but failed to save to database")

            except Exception as e:
                st.error(f"❌ Draft generation failed: {str(e)}")
                draft = create_fallback_draft(case_details, st.session_state.medical_text)
                st.session_state.current_draft = draft

    # Handle AI Rewrite Trigger
    if st.session_state.ai_rewrite_trigger:
        with st.spinner("🔄 Rewriting draft with AI..."):
            current_draft_content = st.session_state.get("current_draft", "")
            improvement_goal = st.session_state.ai_rewrite_goal
            
            if current_draft_content.strip() and improvement_goal:
                improved_draft = rewrite_draft_with_ai(current_draft_content, improvement_goal)
                if improved_draft and "Error" not in improved_draft:
                    st.session_state.current_draft = improved_draft
                    st.session_state.ai_rewrite_trigger = False
                    st.session_state.ai_rewrite_goal = ""
                    st.success("✅ Draft improved successfully!")
                else:
                    st.error("❌ Failed to rewrite draft. Please try again.")
                    st.session_state.ai_rewrite_trigger = False
            else:
                st.error("No draft content or improvement goal provided.")
                st.session_state.ai_rewrite_trigger = False

    # Display Generated Draft
    if st.session_state.get("current_draft"):
        st.markdown("---")
        st.subheader("📄 Generated Draft")
        
        # Editable draft area with better styling
        edited_draft = st.text_area(
            "**Review and Edit Draft:**",
            value=st.session_state.current_draft,
            height=400,
            key="draft_editor"
        )
        
        # Update current draft if edited
        if edited_draft != st.session_state.current_draft:
            st.session_state.current_draft = edited_draft

        # Export and Rewrite Options
        st.subheader("🛠️ Document Tools")
        
        # Export buttons in a nice layout
        col1, col2 = st.columns(2)
        
        with col1:
            if st.button("📄 **Export as Word**", use_container_width=True, key="word_export_btn"):
                current_draft = st.session_state.get("current_draft", "")
                if not current_draft.strip():
                    st.error("❌ No draft content available to export.")
                else:
                    with st.spinner("🔄 Generating Word document..."):
                        word_buffer = create_professional_word_download(current_draft, "legal_demand_draft")
                        # Create unique key for download button
                        download_key = f"word_download_{datetime.now().strftime('%H%M%S')}"
                        st.download_button(
                            label="⬇️ **Download Word Document**",
                            data=word_buffer,
                            file_name="legal_demand_draft.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key=download_key
                        )
                        st.success("✅ Word document generated successfully!")
                    
        with col2:
            if st.button("📊 **Export as PDF**", use_container_width=True, key="pdf_export_btn"):
                current_draft = st.session_state.get("current_draft", "")
                if not current_draft.strip():
                    st.error("❌ No draft content available to export.")
                else:
                    with st.spinner("🔄 Generating PDF..."):
                        pdf_buffer = create_professional_pdf_download(current_draft, "legal_demand_draft")
                        
                        if pdf_buffer:
                            # Check if buffer has content
                            pdf_buffer.seek(0, 2)
                            size = pdf_buffer.tell()
                            pdf_buffer.seek(0)
                            
                            if size > 100:
                                download_key = f"pdf_download_{datetime.now().strftime('%H%M%S')}"
                                st.download_button(
                                    label="⬇️ **Download PDF Document**",
                                    data=pdf_buffer,
                                    file_name="legal_demand_draft.pdf",
                                    mime="application/pdf",
                                    use_container_width=True,
                                    key=download_key
                                )
                                st.success("✅ PDF generated successfully!")
                            else:
                                st.error("❌ Generated PDF is empty. Please try again.")
                        else:
                            st.error("❌ PDF generation failed. Please try again.")

        # AI Rewrite Section
        st.subheader("✨ AI Rewrite Options")
        col1, col2 = st.columns([3, 1])
        with col1:
            rewrite_goal = st.selectbox(
                "**Choose improvement goal:**",
                st.session_state.rewrite_goals,
                key="rewrite_goal_select"
            )
            custom_goal = st.text_input(
                "**Or enter custom goal:**",
                placeholder="e.g., Make tone more persuasive, improve clarity...",
                key="custom_goal_input"
            )
        with col2:
            st.markdown("<br>", unsafe_allow_html=True)
            if st.button("🔄 **AI Rewrite**", use_container_width=True, key="ai_rewrite_btn"):
                final_goal = custom_goal.strip() if custom_goal.strip() else rewrite_goal
                if not final_goal:
                    st.warning("⚠️ Please select or enter an improvement goal.")
                else:
                    # Set trigger for AI rewrite
                    st.session_state.ai_rewrite_trigger = True
                    st.session_state.ai_rewrite_goal = final_goal
                    st.rerun()

# Sidebar setup
def setup_sidebar():
    """Setup sidebar content with working logout"""
    st.sidebar.title(f"Welcome {st.session_state.name}!")
    st.sidebar.write(f"Role: {st.session_state.user_role}")
    st.sidebar.markdown("---")
    
    # AI System Status
    ai_working, ai_msg = check_ai_setup()
    if ai_working:
        st.sidebar.success(ai_msg)
    else:
        st.sidebar.error(ai_msg)
    
    # Database Stats
    if st.session_state.user_role == "admin":
        draft_stats = get_all_drafts_stats()
        st.sidebar.markdown("---")
        st.sidebar.subheader("📊 System Stats")
        st.sidebar.write(f"Total Drafts: **{draft_stats['total_drafts']}**")
        st.sidebar.write(f"Active Users: **{len(get_all_users())}**")
    
    # User Stats
    user_draft_count = get_user_drafts_count(st.session_state.user_id)
    st.sidebar.markdown("---")
    st.sidebar.subheader("Your Stats")
    st.sidebar.write(f"📝 Your Drafts: **{user_draft_count}**")
    
    # Logout button with direct action
    if st.sidebar.button("🚪 Logout", key="sidebar_logout"):
        # Clear all session state
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()

# Main application
def main():
    # Initialize authentication and session state
    initialize_authentication()
    initialize_session_state()

    # Check authentication
    if not check_authentication():
        col1, col2, col3 = st.columns([1,2,1])
        with col2:
            st.title("⚖️ Legal Draft Assistant")
            login()
        return

    # Setup sidebar
    setup_sidebar()

    # Route to appropriate dashboard
    if st.session_state.user_role == "admin":
        show_admin_dashboard()
    else:
        show_staff_dashboard()

if __name__ == "__main__":
    main()
